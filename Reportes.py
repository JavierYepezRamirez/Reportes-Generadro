# Generador de resporte 
# Funciona unicamente con "Reportes CINERGIA" y "NODOS"
# Es un programa que hace una copia del Word original y edita los campos seleccionado,
# tambien permite agregar actividades mas de las que ya estan asignadas por defecto,
# la cuales deben de estar en un json llamado "actividades.json", el programa permite 
# que el usuario pueda seleccionar fotos guardadas en su equipo y pueda ver cuales son
# mostrando una vista previa y la eliminsacion de las imagenes que no se necesiten
# @Version 1.1 Mejorado de vista, problemas con el scroll solucionados, 
# problemas las miniaturas se agregan cada 3, mejorado de mostrado de horas, pantalla completa y inplementacion de logo.
# 26/06/2025
# By: Javier Yepez Ramirez

# *Cosas a agregar*
#   - Fondo

# Forma de enpaquetado
# pyinstaller --noconfirm --onefile --windowed --add-data "icono.ico;." --add-data "actividades.json;." --add-data "REPORTE CINERIA.docx;." --icon=icono.ico --distpath "D:\Reportes" Reportes.py

import os
import sys
import shutil
import datetime
import tkinter as tk
from tkinter import ttk, filedialog, messagebox
from tkcalendar import DateEntry
from openpyxl import load_workbook
from docx import Document
from docx.shared import Inches
import json
from PIL import Image, ImageTk

root = tk.Tk()

# Ruta al icono
icon_path = os.path.join(os.path.dirname(__file__), 'icono.ico')
print(f"Ruta del icono: {icon_path}")

if os.path.exists(icon_path):
    try:
        root.iconbitmap(icon_path)
    except Exception as e:
        print(f"No se pudo cargar el icono: {e}")
else:
    print("El archivo icono.ico no fue encontrado.")

# Función para recursos empaquetados con PyInstaller
def get_path(relative_path):
    if hasattr(sys, '_MEIPASS'):
        return os.path.join(sys._MEIPASS, relative_path)
    return os.path.join(os.path.abspath("."), relative_path)

# Archivos empaquetados
reporte_origen = get_path("REPORTE CINERIA.docx")
actividades_path = get_path("actividades.json")
logo_path = get_path("logo.png")
fondo_path = get_path("fondo.png")

# Archivos externos (NO empaquetados)
excel_path = os.path.join(os.getcwd(), "NODOS.xlsx")
reporte_destino = os.path.join(os.getcwd(), "temp_reporte.docx")

# Abrir Excel
wb = load_workbook(excel_path)
ws = wb.active


clientes_data = []
for row in ws.iter_rows(min_row=2, values_only=True):
    clientes_data.append({
        "nombre_nodo": row[3],
        "municipio": row[13],
        "direccion": row[17],
        "codigo_postal": row[18],
        "nombre_sitio": row[3],
        "tipo_espacio": row[4],
        "entidad": row[6],
        "latitud": row[14],
        "longitud": row[15],
        "clave_nodo": row[16],
        "nomenclatura_redgto": row[2]
    })

nombres_nodo = [c["nombre_nodo"] for c in clientes_data]
imagenes = []

# --- Funciones --- 
def reemplazar_texto(doc, buscar, reemplazo):
    reemplazo = str(reemplazo)
    
    # Reemplazo en párrafos normales
    for p in doc.paragraphs:
        if buscar in p.text:
            for run in p.runs:
                if run.text and buscar in run.text:
                    run.text = run.text.replace(buscar, reemplazo)
    
    # Reemplazo en tablas
    for tabla in doc.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                for p in celda.paragraphs:
                    if buscar in p.text:
                        for run in p.runs:
                            if run.text and buscar in run.text:
                                run.text = run.text.replace(buscar, reemplazo)

    # Reemplazo en encabezado (header)
    for section in doc.sections:
        header = section.header
        for p in header.paragraphs:
            if buscar in p.text:
                for run in p.runs:
                    if run.text and buscar in run.text:
                        run.text = run.text.replace(buscar, reemplazo)

def generar_reporte():
    seleccionado = nodo_var.get()
    datos = next((c for c in clientes_data if c["nombre_nodo"] == seleccionado), None)
    if not datos:
        messagebox.showerror("Error", "Nodo no encontrado.")
        return
    try:
        shutil.copy(reporte_origen, reporte_destino)
    except FileNotFoundError:
        messagebox.showerror("Error", "No se encontró 'REPORTE CINERIA.docx'")
        return

    doc = Document(reporte_destino)
    reemplazos = {
        "cliente": "RED GUANAJUATO",
        "municipio": datos['municipio'],
        "direccion": datos['direccion'],
        "codigo": datos['codigo_postal'],
        "nombre": f"{datos['clave_nodo']} - {datos['nombre_sitio']}",
        "tipo de espacio": datos['tipo_espacio'],
        "entidad": entidad_var.get(),
        "latitud": str(datos['latitud']),
        "longitud": str(datos['longitud']),
        "id": datos['nomenclatura_redgto'],
        "emision": fecha_emision.get(),
        "fecha de apertura": fecha_apertura.get(),
        "llegada": fecha_llegada.get(),
        "cierre": fecha_cierre.get(),
        "Trabajador": trabajador_var.get(),
        "Hora": hora_var.get(),
        "Tecnico": tecnico_var.get(),
        "Servicio1": servicio_var.get(),
        "act": '\n'.join(
            listbox_actividades.get(i) for i in orden_seleccion_actividades
        ) or "",
        "Mantenimiento1": mantenimiento_var.get(),
    }

    for buscar, reemplazo in reemplazos.items():
        reemplazar_texto(doc, buscar, reemplazo)

    if imagenes:
        for tabla in doc.tables:
            for fila in tabla.rows:
                for celda in fila.cells:
                    if "Subir fotos" in celda.text:
                        for img_path in imagenes:
                            try:
                                p = celda.add_paragraph()
                                r = p.add_run()
                                r.add_picture(img_path, width=Inches(3))
                            except Exception as e:
                                messagebox.showerror("Error al insertar imagen", f"{img_path}\n{e}")
                        break

    guardar_como = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Document", "*.docx")])
    if guardar_como:
        doc.save(guardar_como)
        if os.path.exists(reporte_destino):
            os.remove(reporte_destino)
        messagebox.showinfo("Éxito", f"Reporte guardado:\n{guardar_como}")
        limpiar_campos()

def limpiar_campos():
    hoy = datetime.date.today()
    fecha_emision.set_date(hoy)
    fecha_apertura.set_date(hoy)
    fecha_llegada.set_date(hoy)
    fecha_cierre.set_date(hoy)

    nodo_var.set('')
    entidad_var.set('')
    trabajador_var.set('')
    hora_var.set('')
    tecnico_var.set('')
    servicio_var.set('')
    nueva_actividad_var.set("")
    mantenimiento_var.set('')
    imagenes.clear()
    lbl_fotos.config(text="0 foto(s) seleccionada(s)")
    orden_seleccion_actividades.clear()


def mostrar_creditos():
    ventana_creditos = tk.Toplevel(root)
    ventana_creditos.title("Créditos")
    ventana_creditos.geometry("400x200")
    ventana_creditos.configure(bg="#f2f2f2")
    
    texto_creditos = (
        "Generador de Reportes CINERGIA\n"
        "Desarrollado por: Javier Yepez Ramírez\n"
        "Universidad de Guanajuato - DICIS\n"
        "2025\n"
        "\n"
        "¡Gracias por usar esta aplicación!"
    )
    
    lbl_creditos = ttk.Label(ventana_creditos, text=texto_creditos, justify="center", background="#f2f2f2", font=("Segoe UI", 11))
    lbl_creditos.pack(expand=True, padx=20, pady=20)
    
    btn_cerrar = ttk.Button(ventana_creditos, text="Cerrar", command=ventana_creditos.destroy)
    btn_cerrar.pack(pady=(0, 2))

def agregar_actividad():
    valor_entrada = entry_nueva_actividad.get().strip()
    if valor_entrada and valor_entrada not in listbox_actividades.get(0, "end"):
        listbox_actividades.insert("end", valor_entrada)
        actividades_lista.append(valor_entrada)
        with open(actividades_path, "w", encoding="utf-8") as f:
            json.dump(actividades_lista, f, ensure_ascii=False, indent=2)
        nueva_actividad_var.set("")
    elif not valor_entrada:
        messagebox.showwarning("Campo vacío", "Escribe una actividad para agregarla.")
    else:
        messagebox.showinfo("Duplicado", "La actividad ya está en la lista.")

class AutocompleteCombobox(ttk.Combobox):
    def set_completion_list(self, completion_list):
        self._completion_list = sorted(completion_list, key=str.lower)
        self['values'] = self._completion_list
        self.bind('<KeyRelease>', self.handle_keyrelease)

    def handle_keyrelease(self, event):
        if event.keysym in ("Up", "Down", "Return", "Escape"):
            # Solo abrir desplegable si usuario presiona flecha abajo
            if event.keysym == "Down":
                self.event_generate('<Down>')
            return

        typed = self.get().lower()
        if typed == "":
            filtered = self._completion_list
        else:
            filtered = [item for item in self._completion_list if typed in item.lower()]

        self['values'] = filtered

        # Mantener el texto que el usuario escribió
        current_text = self.get()
        self.delete(0, tk.END)
        self.insert(0, current_text)


def actualizar_orden_seleccion(event):
    seleccion_actual = listbox_actividades.curselection()
    for i in seleccion_actual:
        if i not in orden_seleccion_actividades:
            orden_seleccion_actividades.append(i)

    # Elimina los índices que ya no están seleccionados
    for i in orden_seleccion_actividades[:]:
        if i not in seleccion_actual:
            orden_seleccion_actividades.remove(i)
            
def mostrar_miniaturas():
    # Limpiar el frame
    for widget in frame_miniaturas.winfo_children():
        widget.destroy()

    thumbnails.clear()
    for idx, ruta in enumerate(imagenes):
        try:
            img = Image.open(ruta)
            img.thumbnail((280, 280))
            img_tk = ImageTk.PhotoImage(img)
            thumbnails.append(img_tk)  # guardar referencia

            marco = tk.Frame(frame_miniaturas, bd=1, relief="raised")
            marco.grid(row=idx // 3, column=idx % 3, padx=4, pady=4)  # distribuye en filas de 3

            lbl_img = tk.Label(marco, image=img_tk)
            lbl_img.pack()

            btn_x = tk.Button(marco, text="X", command=lambda i=idx: eliminar_imagen(i), bg="#e74c3c", fg="white")
            btn_x.pack(fill="x")

        except Exception as e:
            print(f"Error al cargar imagen: {ruta}", e)

def eliminar_imagen(indice):
    if 0 <= indice < len(imagenes):
        del imagenes[indice]
        lbl_fotos.config(text=f"{len(imagenes)} foto(s) seleccionada(s)")
        mostrar_miniaturas()

def subir_fotos():
    archivos = filedialog.askopenfilenames(filetypes=[("Imágenes", "*.jpg *.jpeg *.png")])
    if archivos:
        imagenes.extend(archivos)
        lbl_fotos.config(text=f"{len(imagenes)} foto(s) seleccionada(s)")
        mostrar_miniaturas()

#Variables globales
orden_seleccion_actividades = []
imagenes = []  
thumbnails = [] 

# --- INTERFAZ ---
root.title("Generador de Reportes CINERGIA")
root.state("zoomed")

style = ttk.Style(root)
style.theme_use('clam')
style.configure("TLabel", background="#FFFFFF", font=("Segoe UI", 12))
style.configure("TCombobox", font=("Segoe UI", 12))
style.configure("TEntry", font=("Segoe UI", 12))
style.configure("TButton", font=("Segoe UI", 12, "bold"), padding=6)
style.map("TButton",
          foreground=[('active', 'white')],
          background=[('active', '#2ecc71')])
style.configure("White.TFrame", background="white")

# Scrollable main_frame dentro de un canvas
frame_canvas = ttk.Frame(root, style="White.TFrame")
frame_canvas.pack(fill="both", expand=True)

canvas = tk.Canvas(frame_canvas, highlightthickness=0, background="white")
scroll_y = ttk.Scrollbar(frame_canvas, orient="vertical", command=canvas.yview)
canvas.configure(yscrollcommand=scroll_y.set)

scroll_y.pack(side="right", fill="y")
canvas.pack(side="left", fill="both", expand=True)

# Frame interno desplazable
scrollable_frame = tk.Frame(canvas, bg="white")
scrollable_window_id = canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")

def _on_mousewheel(event):
    if canvas.winfo_height() < canvas.bbox("all")[3]:  
        if sys.platform == 'darwin':  # macOS
            canvas.yview_scroll(-1 * int(event.delta), "units")
        else:
            canvas.yview_scroll(-1 * int(event.delta / 60), "units")  

def actualizar_scrollregion(event=None):
    canvas.configure(scrollregion=canvas.bbox("all"))

scrollable_frame.bind("<Configure>", actualizar_scrollregion)

# Scroll solo cuando el mouse está sobre el canvas
canvas.bind("<Enter>", lambda e: canvas.bind_all("<MouseWheel>", _on_mousewheel))
canvas.bind("<Leave>", lambda e: canvas.unbind_all("<MouseWheel>"))

main_frame = scrollable_frame

# Carga el logo y redimensiónalo si quieres (opcional)
logo_img_orig = Image.open(logo_path)
logo_img_resized = logo_img_orig.resize((150, 150), Image.Resampling.LANCZOS)  # Ajusta tamaño
logo_img_tk = ImageTk.PhotoImage(logo_img_resized)

# Label con el logo
logo_label = ttk.Label(main_frame, image=logo_img_tk, background="white")
logo_label.grid(row=0, column=0, padx=180, pady=10, sticky="w")

# El título pasa a la columna 1
titulo = ttk.Label(main_frame, text="Generador de Reportes CINERGIA", font=("Arial", 30, "bold"), background="white")
titulo.grid(row=0, column=1, pady=(0, 20), sticky="w")

nodo_var = tk.StringVar()
trabajador_var = tk.StringVar()
hora_var = tk.StringVar()
tecnico_var = tk.StringVar()
entidad_var = tk.StringVar()
servicio_var = tk.StringVar()
mantenimiento_var = tk.StringVar()
nueva_actividad_var = tk.StringVar()

fecha_emision = DateEntry(main_frame, width=18, date_pattern="dd/mm/yyyy")
fecha_apertura = DateEntry(main_frame, width=18, date_pattern="dd/mm/yyyy")
fecha_llegada = DateEntry(main_frame, width=18, date_pattern="dd/mm/yyyy")
fecha_cierre = DateEntry(main_frame, width=18, date_pattern="dd/mm/yyyy")

# Cargar actividades guardadas
if os.path.exists(actividades_path):
    with open(actividades_path, "r", encoding="utf-8") as f:
        actividades_lista = json.load(f)
else:
    actividades_lista = []

    # Asegúrate que nombres_nodo no contenga None
nombres_nodo_filtrados = [n for n in nombres_nodo if n is not None]

# Crea el combobox autocompletable
combo_nodo = AutocompleteCombobox(main_frame, textvariable=nodo_var, width=40)
combo_nodo.set_completion_list(nombres_nodo_filtrados)

etiquetas = [
   ("Selecciona un nodo:", combo_nodo),
    ("Fecha de emisión:", fecha_emision),
    ("Fecha de apertura:", fecha_apertura),
    ("Fecha de llegada al sitio:", fecha_llegada),
    ("Fecha de cierre:", fecha_cierre),
    ("Trabajador:", ttk.Combobox(main_frame, textvariable=trabajador_var, values=["Jaime López Horta", "Juan Manuel Manríquez Sarabia", "Ricardo Garcidueñas Vargas", "Otro"], width=30)),
    ("Hora:", ttk.Combobox(main_frame, textvariable=hora_var, 
                       values=[f"{h}:{m:02} {'AM' if h < 12 else 'PM'}" 
                               for h in range(8, 17) 
                               for m in (0, 30)], 
                       width=20)),
    ("Área de técnico:", ttk.Combobox(main_frame, textvariable=tecnico_var, values=["Jardín", "Kiosco", "Presidencia ", "Dirección ", "Patio ", "Aula ", "Biblioteca"], width=30)),
    ("Entidad:", ttk.Combobox(main_frame, textvariable=entidad_var, values=["Publico", "Preescolar", "Primaria", "Telesecundaria", "SABES", "UVEG"], width=30)),
    ("Tipo de servicio:", ttk.Combobox(main_frame, textvariable=servicio_var, values=["Servicio Preventivo", "Servicio Correctivo"], width=40)),
    ("Tipo de mantenimiento:", ttk.Combobox(main_frame, textvariable=mantenimiento_var, values=["Mantenimiento Preventivo", "Mantenimiento Correctivo"], width=40))
]

for i, (texto, widget) in enumerate(etiquetas):
    lbl = ttk.Label(main_frame, text=texto)
    lbl.grid(row=i + 1, column=0, sticky="e", pady=4, padx=5)
    widget.grid(row=i + 1, column=1, sticky="ew", pady=4)
    
# Actividades
fila_actividades = len(etiquetas) + 1
lbl_actividades = ttk.Label(main_frame, text="Actividades realizadas:")
lbl_actividades.grid(row=fila_actividades, column=0, sticky="ne", pady=(10, 0), padx=5)

# Frame contenedor para el Listbox y su scrollbar
frame_actividades = tk.Frame(main_frame)
frame_actividades.grid(row=fila_actividades, column=1, sticky="w", pady=(10, 0))

# Scrollbar vertical para el listbox
scrollbar_actividades = ttk.Scrollbar(frame_actividades, orient="vertical")
scrollbar_actividades.pack(side="right", fill="y")

# Listbox con scrollbar asociado
listbox_actividades = tk.Listbox(frame_actividades, selectmode="multiple", height=20, width=75, exportselection=False, yscrollcommand=scrollbar_actividades.set)
for act in actividades_lista:
    listbox_actividades.insert("end", act)
listbox_actividades.pack(side="left", fill="both")

scrollbar_actividades.config(command=listbox_actividades.yview)

# Agregar actividad
fila_entrada = fila_actividades + 3

# Entry para nueva actividad, con un ancho decente
entry_nueva_actividad = ttk.Entry(main_frame, textvariable=nueva_actividad_var, width=80)
entry_nueva_actividad.grid(row=fila_entrada, column=1, sticky="w", pady=(4, 0), padx=(0, 100))

# Botón separado con padx para espacio
btn_agregar_actividad = ttk.Button(main_frame, text="Agregar actividad", command=lambda: agregar_actividad())
btn_agregar_actividad.grid(row=fila_entrada, column=1, sticky="w", pady=(4, 0), padx=(510, 0))

# Vinculación
listbox_actividades.bind('<<ListboxSelect>>', actualizar_orden_seleccion)

# Botón subir fotos y etiqueta
fila_fotos = fila_entrada + 1
btn_subir_fotos = ttk.Button(main_frame, text="Subir fotos", command=subir_fotos)
btn_subir_fotos.grid(row=fila_fotos, column=0, sticky="w", pady=12, padx=5)

lbl_fotos = ttk.Label(main_frame, text="0 foto(s) seleccionada(s)")
lbl_fotos.grid(row=fila_fotos, column=1, sticky="w")

frame_miniaturas = tk.Frame(main_frame, bg="white")
frame_miniaturas.grid(row=fila_fotos + 1, column=0, columnspan=2, pady=(5, 0), sticky="w")

# Botón generar
fila_generar = fila_fotos + 3
btn_generar = ttk.Button(main_frame, text="Generar reporte", command=generar_reporte)
btn_generar.grid(row=fila_generar, column=0, pady=20, padx=5, sticky="w")

# Botón créditos
btn_creditos = ttk.Button(root, text="?", width=3, command=mostrar_creditos, style="TButton")
btn_creditos.place(relx=0.98, rely=0.02, anchor="ne")

limpiar_campos()

root.mainloop()

